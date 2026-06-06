from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x5E, 0x97)

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shd)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(10.5)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        fill = 'FFFFFF' if ri % 2 == 0 else 'EBF3FB'
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            tc = cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_heading('🛡️ Cybersecurity Incident Management System (CIMS)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Detailed Project Report\nArchitecture, Core Functionalities, and Database Schema')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].italic = True
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 – SPECIFICATIONS & CORE FUNCTIONALITIES
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 1: System Specifications & Core Functionalities')
para('CIMS is a robust, multi-tenant platform designed for Security Operations Centers (SOC) to track, manage, and resolve cybersecurity threats across multiple organizations and infrastructure assets.')

heading2('System Specifications')
bullet('Frontend: React.js, Vite, Axios, React Router, Lucide React (Icons).')
bullet('Backend API: Node.js, Express.js, JSON Web Tokens (JWT) for Auth.')
bullet('Database: PostgreSQL (Relational schema with Advanced Triggers and Views).')
bullet('Security: bcrypt password hashing, Role-Based Access Control (RBAC), API Rate Limiting, CORS protection.')

heading2('Core Functionalities')
bullet('Incident Lifecycle Management: Create incidents (atomic DB transactions), transition statuses securely (Open -> In Progress -> Resolved -> Closed), and automatically log remediation timestamps via PostgreSQL triggers.')
bullet('Role-Based Access Control (RBAC): Strict hierarchies (Admin, Manager, Lead, Senior, Junior). For example, only Admins and Managers can create incidents or escalate severity.')
bullet('Vulnerability Tracking: Map known CVEs (Common Vulnerabilities and Exposures) to physical and cloud infrastructure assets with real-time patching statuses.')
bullet('Audit & Remediation Logging: Every state change to an incident is irreversibly logged to a Remediation_Actions table for compliance and forensics.')
bullet('Interactive Dashboard: Analytical views pulling from pre-joined PostgreSQL Views to calculate active threats and analyst workload.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 – BACKEND ARCHITECTURE & FILE DESCRIPTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 2: Backend Architecture & File Descriptions')
para('The backend follows an MVC-like, scalable modular architecture separating routing, middleware logic, and database operations.')

heading2('Core Server & Setup')
add_table(
    ['File', 'Purpose'],
    [
        ['src/server.js', 'Entry point. Validates DB connection, binds to the network port, handles graceful shutdown on SIGTERM/SIGINT.'],
        ['src/app.js', 'Express app setup. Registers CORS, parses JSON, applies global rate limiting, and mounts all API routers.'],
        ['src/config/db.js', 'Initializes the standard pg connection pool using environment variables to connect to PostgreSQL.'],
    ]
)

heading2('Business Logic (Controllers)')
add_table(
    ['File', 'Purpose'],
    [
        ['authController.js', 'Authenticates users by comparing bcrypt password hashes and issues secure JWT tokens.'],
        ['incidentController.js', 'Core logic: Creates incidents (with DB transactions), updates statuses with strict workflow validation, escalates severity, and assigns analysts/assets.'],
        ['assetController.js', 'Manages physical/cloud assets. Links assets to vulnerabilities (CVEs) using ON CONFLICT UPSERT logic.'],
        ['reportController.js', 'Executes complex aggregated SQL queries to generate dashboard statistics (e.g., analyst workload, critical unpatched assets).'],
        ['organizationController.js', 'CRUD operations for tenant organizations.'],
        ['analystController.js', 'Manages analyst profiles, soft-deletes (is_active), and assigns RBAC roles.'],
    ]
)

heading2('Security & Middleware')
add_table(
    ['File', 'Purpose'],
    [
        ['auth.js', 'Intercepts requests to verify JWT signatures and extract analyst metadata (user ID, role).'],
        ['authorize.js', 'RBAC middleware. Blocks requests if the user lacks the required role hierarchy (returns 403 Forbidden).'],
        ['validate.js', 'Request body validation. Ensures required fields and valid enums are passed before hitting the DB.'],
        ['rateLimit.js', 'Uses express-rate-limit to protect endpoints from brute-force and DoS attacks.'],
        ['errorHandler.js', 'Centralized error catching mapping ApiError instances to clean JSON responses.'],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 – SCHEMA OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 3: Database Schema — All 9 Tables')
para('The database is organized into 3 tiers:')
add_table(
    ['Tier', 'Tables'],
    [
        ['Core Entities', 'Organizations, Analysts, Assets, Vulnerabilities'],
        ['Junction / Bridge', 'Asset_Vulnerabilities, Incident_Assets, Incident_Analysts'],
        ['Operational', 'Incidents, Remediation_Actions'],
    ]
)

heading2('TABLE 1 — Organizations')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['org_id','SERIAL','PRIMARY KEY','Auto-incremented unique ID'],
        ['name','VARCHAR(255)','NOT NULL','Organization name'],
        ['contact_email','VARCHAR(255)','UNIQUE','One email per org'],
        ['created_at','TIMESTAMP','DEFAULT NOW()','Record creation time'],
    ]
)

heading2('TABLE 2 — Analysts')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['analyst_id','SERIAL','PRIMARY KEY','Unique analyst ID'],
        ['org_id','INT','FK → Organizations','Which org they belong to'],
        ['email','VARCHAR(255)','NOT NULL, UNIQUE','Login email'],
        ['role','VARCHAR(50)','CHECK (in list)','Junior/Senior/Lead/Manager/Admin'],
        ['password_hash','VARCHAR(255)','NOT NULL','bcrypt hashed password'],
        ['is_active','BOOLEAN','DEFAULT TRUE','Soft-delete flag'],
    ]
)

heading2('TABLE 3 — Assets')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['asset_id','SERIAL','PRIMARY KEY','Unique asset ID'],
        ['organization_id','INT','FK → Organizations','Owner organization'],
        ['type','VARCHAR(100)','CHECK (in list)','Server/Router/Firewall/etc.'],
        ['ip_address','INET','NOT NULL, UNIQUE','PostgreSQL native IP validation'],
        ['criticality','VARCHAR(50)','CHECK (in list)','Low/Medium/High/Critical'],
    ]
)

heading2('TABLE 4 — Vulnerabilities')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['id','SERIAL','PRIMARY KEY','Auto ID'],
        ['cve_id','VARCHAR(50)','NOT NULL, UNIQUE','CVE number (e.g., CVE-2023-1234)'],
        ['cvss_score','DECIMAL(3,1)','CHECK (0.0–10.0)','CVSS severity score'],
    ]
)

heading2('TABLE 5 — Asset_Vulnerabilities (Junction Table)')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['asset_id','INT','FK → Assets','Which asset'],
        ['vulnerability_id','INT','FK → Vulnerabilities','Which vulnerability'],
        ['patch_status','VARCHAR(50)','CHECK (in list)','Patched/Unpatched/Excluded/In Progress'],
    ]
)
bullet('Composite PK: PRIMARY KEY(asset_id, vulnerability_id) — prevents duplicate mapping.')

heading2('TABLE 6 — Incidents')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['id','SERIAL','PRIMARY KEY','Incident ID'],
        ['severity','VARCHAR(50)','CHECK + NOT NULL','Low/Medium/High/Critical'],
        ['status','VARCHAR(50)','CHECK, DEFAULT Open','Open/In Progress/On Hold/Resolved/Closed'],
        ['created_by','INT','FK → Analysts, SET NULL','Who filed the incident'],
        ['resolved_at','TIMESTAMP','—','Auto-set by trigger when Resolved'],
    ]
)

heading2('TABLE 7 & 8 — Junction Tables (Incident_Assets & Incident_Analysts)')
para('These tables map the many-to-many relationships between an incident and the assets it affects, as well as the incident and the analysts actively investigating it. Both use composite primary keys and ON DELETE CASCADE constraints.')

heading2('TABLE 9 — Remediation_Actions')
add_table(
    ['Column','Type','Constraint','Purpose'],
    [
        ['incident_id','INT','FK → Incidents, CASCADE','Which incident'],
        ['analyst_id','INT','FK → Analysts, SET NULL','Who acted'],
        ['action_taken','TEXT','NOT NULL','Description of action'],
        ['result','VARCHAR(50)','CHECK (in list)','Pending/In Progress/Successful/Partial/Failed'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 – RELATIONSHIPS
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 4: Relationships Between Tables')
add_table(
    ['Relationship','Type','Bridge Table'],
    [
        ['Organization → Analysts','One-to-Many','—'],
        ['Organization → Assets','One-to-Many','—'],
        ['Organization → Incidents','One-to-Many','—'],
        ['Incidents ↔ Analysts','Many-to-Many','Incident_Analysts'],
        ['Incidents ↔ Assets','Many-to-Many','Incident_Assets'],
        ['Assets ↔ Vulnerabilities','Many-to-Many','Asset_Vulnerabilities'],
        ['Incidents → Remediation_Actions','One-to-Many','—'],
        ['Analysts → Remediation_Actions','One-to-Many','—'],
    ]
)

heading2('Referential Integrity (ON DELETE)')
add_table(
    ['Behaviour','Tables Using It','Effect'],
    [
        ['ON DELETE CASCADE','Organizations→Analysts, Assets, Incidents','Deleting parent deletes all related children tables. Secures data cleanup.'],
        ['ON DELETE SET NULL','Incidents.created_by, Remediation.analyst_id','Deleting an analyst NULLs the FK. Prevents destroying vital incident history.'],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 5 – FUNCTIONS / TRIGGERS / VIEWS
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 5: Advanced PostgreSQL Features')

heading2('Function: set_resolved_timestamp() & Trigger')
para('Auto-sets the resolved_at timestamp when an incident status changes to Resolved directly within the database engine.')
code_block(
"""CREATE OR REPLACE FUNCTION set_resolved_timestamp() RETURNS TRIGGER AS $$
BEGIN
    IF NEW.status = 'Resolved' AND OLD.status != 'Resolved' THEN
        NEW.resolved_at = CURRENT_TIMESTAMP;
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trigger_set_resolved_timestamp
BEFORE UPDATE ON Incidents
FOR EACH ROW EXECUTE FUNCTION set_resolved_timestamp();"""
)

heading2('View: Active_Incidents_Dashboard')
para('Provides a pre-compiled, highly optimized snapshot of active threats across organizations without writing complex application logic.')
code_block(
"""CREATE OR REPLACE VIEW Active_Incidents_Dashboard AS
SELECT
    i.id AS incident_id, o.name AS organization_name,
    i.title, i.type, i.severity, i.status, i.created_at,
    COUNT(ia.asset_id) AS affected_assets_count
FROM Incidents i
JOIN Organizations o ON i.organization_id = o.org_id
LEFT JOIN Incident_Assets ia ON i.id = ia.incident_id
WHERE i.status IN ('Open', 'In Progress')
GROUP BY i.id, o.name;"""
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 6 – SQL JOINS IN BACKEND
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 6: SQL JOINs Used in Backend Controllers')

heading2('JOIN 1 — Get Assets + Org Name (assetController.js)')
code_block(
"""SELECT a.*, o.name as organization_name
FROM assets a
LEFT JOIN organizations o ON a.organization_id = o.org_id
ORDER BY a.asset_id"""
)
bullet('Type: LEFT JOIN — Returns all assets, mapping the organization name if available.')

heading2('JOIN 2 — Analyst Workload Report (reportController.js)')
code_block(
"""SELECT a.analyst_id, a.name AS analyst_name,
       COUNT(DISTINCT ia.incident_id)::int AS active_incidents
FROM analysts a
LEFT JOIN incident_analysts ia ON ia.analyst_id = a.analyst_id
LEFT JOIN incidents i ON i.id = ia.incident_id AND i.status IN ('Open', 'In Progress')
GROUP BY a.analyst_id, a.name
ORDER BY active_incidents DESC"""
)
bullet('Type: Two LEFT JOINs — Shows all analysts even those with zero active incidents, aggregating their current active workload.')

heading2('JOIN 3 — Critical Unpatched CVEs (assetController.js)')
code_block(
"""SELECT a.asset_id, a.type, a.ip_address,
       v.cve_id, v.cvss_score, av.patch_status
FROM asset_vulnerabilities av
JOIN assets a ON a.asset_id = av.asset_id
JOIN vulnerabilities v ON v.id = av.vulnerability_id
WHERE av.patch_status = 'Unpatched' AND v.cvss_score >= 9.0
ORDER BY v.cvss_score DESC"""
)
bullet('Type: Two INNER JOINs — Filters assets possessing highly critical (CVSS ≥ 9.0), unpatched vulnerabilities.')

heading2('JOIN 4 — Incident Timeline with UNION ALL (reportController.js)')
code_block(
"""SELECT 'incident_created' AS event_type, i.created_at AS event_time,
       CONCAT('Incident created with status ', i.status) AS details
FROM incidents i WHERE i.id = $1
UNION ALL
SELECT 'remediation_action', ra.created_at, ra.action_taken
FROM remediation_actions ra WHERE ra.incident_id = $1
ORDER BY event_time ASC"""
)
bullet('UNION ALL — Merges incident metadata logs and remediation actions into a single chronological timeline.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 7 – BACKEND → FRONTEND FLOW
# ══════════════════════════════════════════════════════════════════════════════
heading1('Part 7: Backend → Frontend Data Flow')

heading2('Architecture Pipeline')
code_block(
"""PostgreSQL Database (Port 5432)
       │  (SQL Queries via 'pg' connection pool)
       ▼
Node.js / Express Backend (Port 5000)
       │  (REST API with JWT Auth validation)
       ▼
Axios HTTP Client (frontend/src/api/client.js)
       │  (Intercepts requests to attach Bearer Tokens)
       ▼
React + Vite Frontend (Port 3000)
       │  (State managed via useState/useEffect)
       ▼
User's Browser"""
)

heading2('API Route Mapping')
add_table(
    ['Frontend Page','API Endpoint','Controller Method','Key Operations'],
    [
        ['IncidentsPage.jsx','GET /incidents','listIncidents()','Dynamic WHERE clauses filtering by status/severity.'],
        ['IncidentsPage.jsx','POST /incidents','createIncident()','Atomic Database Transaction (BEGIN/COMMIT).'],
        ['IncidentsPage.jsx','PATCH /incidents/:id/status','updateIncidentStatus()','Validates workflow (Open->In Progress->Resolved).'],
        ['AssetsPage.jsx','GET /assets','getAllAssets()','Retrieves infrastructure via LEFT JOIN.'],
        ['ThreatMapPage.jsx','GET /assets/critical','getUnpatchedCritical()','JOINs assets with CVSS≥9.0 vulnerabilities.'],
        ['LoginPage.jsx','POST /auth/login','login()','Validates bcrypt hash, issues 12h expiration JWT.'],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\Administrator\Desktop\CIMS_Detailed_Report.docx'
doc.save(output_path)
print("Detailed Report file saved: " + output_path)
